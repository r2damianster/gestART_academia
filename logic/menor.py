import pandas as pd
from docx import Document
from io import BytesIO
import os
import random


class ReportGeneratorLogic:
    TEMPLATE_FILE_NAME = 'NotasMenores.docx'

    COLUMNAS_POR_PARCIAL = {
        '1': {
            'C1': 'Total P1 - Actuación (Actividades de docencia) (C1) (Real)',
            'C2': 'Total P1 - Producción (Trabajo Autónomo) (C2) (Real)',
            'C3': 'Total P1 - Producción (Práctica y experimentación de aprendizajes) (C3) (Real)',
            'C4': 'Total P1 - Acreditación (Evaluación Final) (C4) (Real)',
            'TOTAL': 'Total P1 (Real)',
            'UMBRAL': 7.0,
            'TEXTO': 'PRIMER'
        },
        '2': {
            'C1': 'Total P2 - Actuación (Actividades de docencia) (C1) (Real)',
            'C2': 'Total P2 - Producción (Trabajo Autónomo) (C2) (Real)',
            'C3': 'Total P2 - Producción (Práctica y experimentación de aprendizajes) (C3) (Real)',
            'C4': 'Total P2 - Acreditación (Evaluación Final) (C4) (Real)',
            'TOTAL': 'Total del curso (Real)',
            'UMBRAL': 14.0,
            'TEXTO': 'SEGUNDO'
        }
    }

    PLACEHOLDER_MAP = {
        "fecha": "fecha",
        "titulo_academico_destinatario": "titulo_destinatario",
        "nombres_apellidos_destinatario": "nombres_destinatario",
        "facultad_extension_destinatario": "facultad_extension",
        "titulo_academico_emisor": "titulo_emisor",
        "nombres_apellidos_emisor": "nombres_emisor",
        "titulo_academico_cc": "titulo_cc",
        "nombres_apellidos_cc": "nombres_cc",
        "asignatura": "asignatura"
    }

    def get_template_path(self):
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        return os.path.join(base_dir, "resources", self.TEMPLATE_FILE_NAME)

    def _replace_placeholders_in_paragraph(self, paragraph, replacements):
        full_text = paragraph.text
        for old_text, new_text in replacements.items():
            if old_text in full_text:
                full_text = full_text.replace(old_text, str(new_text))
                while len(paragraph.runs) > 0:
                    paragraph._element.remove(paragraph.runs[0]._element)
                paragraph.add_run(full_text)

    def _extract_subject_from_filename(self, filename):
        parts = filename.split('--')
        return parts[1].strip() if len(parts) > 2 else "ASIGNATURA DESCONOCIDA"

    def process_excel_data(self, excel_file, parcial_id):
        df = pd.read_excel(excel_file, dtype={'Nombre': str, 'Apellido(s)': str})

        p = self.COLUMNAS_POR_PARCIAL.get(parcial_id)
        if not p:
            raise Exception("Parcial no válido.")

        cols_req = ['Nombre', 'Apellido(s)', p['C1'], p['C2'], p['C3'], p['C4'], p['TOTAL']]
        for col in cols_req:
            if col not in df.columns:
                raise Exception(f"Falta columna: {col}")
            if col not in ['Nombre', 'Apellido(s)']:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)

        df_reprobados = df[df[p['TOTAL']] < p['UMBRAL']].copy()

        estudiantes = []
        for _, row in df_reprobados.iterrows():
            nombre = f"{row['Nombre']} {row['Apellido(s)']}".strip()

            if parcial_id == '1':
                estrategia = "Acompañamiento pedagógico continuo y tutoría académica preventiva."
                espacio = "Aula de clase / Tutorías individuales."
                semana = str(random.randint(2, 7))
            else:
                estrategia = "Recuperación de aprendizajes mediante tutorías grupales e individuales."
                espacio = "Escenarios de tutoría (Presencial/Virtual)."
                semana = "17"

            estudiantes.append({
                "correlacional": len(estudiantes) + 1,
                "nombre": nombre,
                "calificacion": float(row[p['TOTAL']]),
                "estrategia": estrategia,
                "espacio": espacio,
                "semana": semana
            })

        return estudiantes, p

    def generate_report(self, datos, estudiantes, info_parcial):
        template_path = self.get_template_path()

        if not os.path.exists(template_path):
            raise Exception(f"No se encontró la plantilla en: {template_path}")

        document = Document(template_path)

        replacements = {
            f"{{{{{self.PLACEHOLDER_MAP[k]}}}}}": v
            for k, v in datos.items()
            if k in self.PLACEHOLDER_MAP
        }

        replacements["{{parcial_texto}}"] = info_parcial['TEXTO']
        replacements["{{nota_minima}}"] = f"{info_parcial['UMBRAL']:.2f}"

        for paragraph in document.paragraphs:
            self._replace_placeholders_in_paragraph(paragraph, replacements)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(paragraph, replacements)

        if len(document.tables) >= 1:
            t1 = document.tables[0]
            for est in estudiantes:
                row = t1.add_row().cells
                row[0].text = str(est["correlacional"])
                row[1].text = est["nombre"]
                row[2].text = f"{est['calificacion']:.2f}"

        if len(document.tables) >= 2:
            t2 = document.tables[1]
            for est in estudiantes:
                row = t2.add_row().cells
                row[0].text = str(est["correlacional"])
                row[3].text = est["estrategia"]
                row[4].text = est["espacio"]
                row[5].text = est["semana"]

        output = BytesIO()
        document.save(output)
        output.seek(0)
        return output
