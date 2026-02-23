import os
import io
import pandas as pd
import unicodedata
import datetime
import re
from docx import Document

class ConvocatoriaLogic:
    def __init__(self):
        self.base_dir = os.path.dirname(os.path.abspath(__file__))
        self.resources_path = os.path.normpath(os.path.join(self.base_dir, "..", "resources"))

    def normalizar(self, texto):
        """Elimina acentos y prepara texto para ordenamiento."""
        if not texto: return ""
        texto = str(texto).lower()
        texto = unicodedata.normalize('NFD', texto)
        return ''.join(c for c in texto if unicodedata.category(c) != 'Mn')

    def formatear_fecha_reunion(self, fecha_str):
        """Convierte 2026-02-20 en '20 de febrero de 2026'."""
        try:
            if not fecha_str or " de " in str(fecha_str): 
                return fecha_str
            
            fecha_obj = datetime.datetime.strptime(str(fecha_str), '%Y-%m-%d')
            meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", 
                     "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
            return f"{fecha_obj.day} de {meses[fecha_obj.month - 1]} de {fecha_obj.year}"
        except:
            return fecha_str

    def reemplazar_en_documento(self, doc, reemplazos):
        """Reemplaza etiquetas {{LLAVE}} ignorando mayúsculas/minúsculas."""
        limpios = {}
        for k, v in reemplazos.items():
            key_name = str(k).replace('{', '').replace('}', '').strip()
            val = str(v) if v is not None else ""
            limpios[key_name.upper()] = val

        def realizar_cambio(texto):
            nuevo_texto = texto
            # Busca cualquier cosa dentro de {{ }}
            for match in re.findall(r'\{\{(.*?)\}\}', nuevo_texto):
                tag_interna = match.strip().upper()
                if tag_interna in limpios:
                    nuevo_texto = nuevo_texto.replace("{{" + match + "}}", limpios[tag_interna])
            return nuevo_texto

        # Procesar párrafos
        for p in doc.paragraphs:
            if '{{' in p.text:
                p.text = realizar_cambio(p.text)
        
        # Procesar tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if '{{' in p.text:
                            p.text = realizar_cambio(p.text)

    def insertar_estudiantes_en_todas_las_tablas(self, doc, lista_estudiantes):
        """Llena todas las tablas que contengan el marcador {{nombre}}."""
        for table in doc.tables:
            marcador_encontrado = False
            fila_plantilla = None
            
            for row in table.rows:
                texto_celda = row.cells[2].text.lower() if len(row.cells) > 2 else ""
                if "{{nombre}}" in texto_celda:
                    marcador_encontrado = True
                    fila_plantilla = row
                    break
            
            if marcador_encontrado:
                parent = fila_plantilla._tr.getparent()
                parent.remove(fila_plantilla._tr)
                
                for i, estudiante in enumerate(lista_estudiantes, start=1):
                    new_row = table.add_row()
                    new_row.cells[0].text = str(i)
                    new_row.cells[1].text = "Estudiante"
                    new_row.cells[2].text = estudiante
                    new_row.cells[3].text = ""

    def procesar_excel_estudiantes(self, lista_archivos):
        """Une múltiples excels, limpia y ordena nombres."""
        nombres_consolidados = []
        for excel_file in lista_archivos:
            try:
                df = pd.read_excel(excel_file)
                for _, row in df.iterrows():
                    nombre = str(row.iloc[0]).strip()
                    apellido = str(row.iloc[1]).strip()
                    
                    if nombre and apellido and "nan" not in (nombre.lower() + apellido.lower()):
                        nombre_completo = f"{apellido} {nombre}".upper()
                        nombres_consolidados.append(nombre_completo)
            except Exception as e:
                print(f"Error procesando archivo: {e}")
                continue
        
        return sorted(list(set(nombres_consolidados)), key=self.normalizar)

    def generar_docx(self, tipo, datos, excel_files=None):
        """Punto de entrada principal."""
        filename = "Convocatoria_Docentes.docx" if tipo == "docente" else "Convocatoria_Estudiantes.docx"
        path = os.path.join(self.resources_path, filename)
        
        if not os.path.exists(path):
            raise FileNotFoundError(f"No se encontró la plantilla en: {path}")

        doc = Document(path)
        
        # 1. Normalizar las llaves de 'datos' para que la fecha se procese bien
        # Buscamos 'fecha_reunion' sin importar si el usuario la mandó como {{FECHA_REUNION}}
        datos_normalizados = {}
        for k, v in datos.items():
            key_limpia = k.replace('{','').replace('}','').lower()
            if 'fecha_reunion' in key_limpia:
                v = self.formatear_fecha_reunion(v)
            datos_normalizados[k] = v

        # 2. Reemplazar texto
        self.reemplazar_en_documento(doc, datos_normalizados)

        # 3. Si hay estudiantes, procesar tablas
        if tipo == "estudiante" and excel_files:
            estudiantes = self.procesar_excel_estudiantes(excel_files)
            if estudiantes:
                self.insertar_estudiantes_en_todas_las_tablas(doc, estudiantes)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer