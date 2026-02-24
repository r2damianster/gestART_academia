import os
import io
from docxtpl import DocxTemplate
from docx import Document
from docxcompose.composer import Composer
from datetime import timedelta

def generar_documento_pat04(datos):
    base_path = os.path.dirname(os.path.abspath(__file__))
    plantilla_path = os.path.join(base_path, "..", "..", "resources", "PAT-03-G-001-F-004.docx")

    try:
        h, m = map(int, datos["hora"].split(':'))
        hora_finalizada = f"{h+2:02d}:{m:02d}"
    except:
        hora_finalizada = "18:00"

    fechas = [(datos["fecha_final"] - timedelta(weeks=i)).strftime("%d/%m/%Y") for i in range(8, -1, -1)]

    # Primera sesión
    tpl = DocxTemplate(plantilla_path)
    ctx1 = {
        "No": 1, "Articulo": datos["articulo"], "NOMBRE": datos["nombre"],
        "FECHA": fechas[0], "HORA": datos["hora"], "HoraFin": hora_finalizada,
        "Tema": datos["temas"][0], "Proxima": datos["temas"][1]
    }
    tpl.render(ctx1)
    
    # Lo pasamos a un objeto Document de docx
    main_buffer = io.BytesIO()
    tpl.save(main_buffer)
    main_buffer.seek(0)
    maestro = Document(main_buffer)
    composer = Composer(maestro)

    # Sesiones restantes
    for i in range(1, 9):
        tpl_temp = DocxTemplate(plantilla_path)
        prox = datos["temas"][i+1] if i < 8 else "Revisión integral para entrega final"
        ctx = {
            "No": i + 1, "Articulo": datos["articulo"], "NOMBRE": datos["nombre"],
            "FECHA": fechas[i], "HORA": datos["hora"], "HoraFin": hora_finalizada,
            "Tema": datos["temas"][i], "Proxima": prox
        }
        tpl_temp.render(ctx)
        
        temp_buffer = io.BytesIO()
        tpl_temp.save(temp_buffer)
        temp_buffer.seek(0)
        
        doc_a_unir = Document(temp_buffer)
        doc_a_unir.paragraphs[0].paragraph_format.page_break_before = True
        composer.append(doc_a_unir)

    final_buffer = io.BytesIO()
    composer.save(final_buffer)
    final_buffer.seek(0)
    return final_buffer