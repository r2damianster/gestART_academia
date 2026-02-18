import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request, render_template, send_file
from io import BytesIO
import os
import random
import re
from datetime import datetime

class ReportGeneratorLogic:
    TEMPLATE_FILE_NAME = 'NotasMenores.docx'
    
    COLUMNAS_POR_PARCIAL = {
        '1': {
            'C1': 'Total P1 - Actuación (Actividades de docencia) (C1) (Real)',
            'C2': 'Total P1 - Producción (Trabajo Autónomo) (C2) (Real)',
            'C3': 'Total P1 - Producción (Práctica y experimentación de aprendizajes) (C3) (Real)',
            'C4': 'Total P1 - Acreditación (Evaluación Final) (C4) (Real)',
            'TOTAL': 'Total P1 (Real)',
            'UMBRAL': 7.0,
            'TEXTO': 'PRIMER'
        },
        '2': {
            'C1': 'Total P2 - Actuación (Actividades de docencia) (C1) (Real)',
            'C2': 'Total P2 - Producción (Trabajo Autónomo) (C2) (Real)',
            'C3': 'Total P2 - Producción (Práctica y experimentación de aprendizajes) (C3) (Real)',
            'C4': 'Total P2 - Acreditación (Evaluación Final) (C4) (Real)',
            'TOTAL': 'Total del curso (Real)',
            'UMBRAL': 14.0,
            'TEXTO': 'SEGUNDO'
        }
    }

    PLACEHOLDER_MAP = {
        "fecha": "fecha", "titulo_academico_destinatario": "titulo_destinatario",
        "nombres_apellidos_destinatario": "nombres_destinatario",
        "facultad_extension_destinatario": "facultad_extension",
        "titulo_academico_emisor": "titulo_emisor", "nombres_apellidos_emisor": "nombres_emisor",
        "titulo_academico_cc": "titulo_cc", "nombres_apellidos_cc": "nombres_cc", "asignatura": "asignatura"
    }

    def get_template_path(self):
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), self.TEMPLATE_FILE_NAME)

    def _replace_placeholders_in_paragraph(self, paragraph, replacements):
        full_text = paragraph.text
        for old_text, new_text in replacements.items():
            if old_text in full_text:
                full_text = full_text.replace(old_text, str(new_text))
                while len(paragraph.runs) > 0:
                    paragraph._element.remove(paragraph.runs[0]._element)
                paragraph.add_run(full_text)

    def _extract_subject_from_filename(self, filename):
        parts = filename.split('--')
        return parts[1].strip() if len(parts) > 2 else "ASIGNATURA DESCONOCIDA"

    def _process_excel_data(self, excel_file, parcial_id):
        try:
            # Forzamos que Nombre y Apellido se lean como texto para evitar el 0.0
            df = pd.read_excel(excel_file, dtype={'Nombre': str, 'Apellido(s)': str})
        except Exception as e:
            raise Exception(f"Error al leer Excel: {e}")

        p = self.COLUMNAS_POR_PARCIAL.get(parcial_id)
        if not p:
            raise Exception("Parcial no válido.")

        # Validar columnas
        cols_req = ['Nombre', 'Apellido(s)', p['C1'], p['C2'], p['C3'], p['C4'], p['TOTAL']]
        for col in cols_req:
            if col not in df.columns:
                raise Exception(f"Falta columna: {col}")
            if col not in ['Nombre', 'Apellido(s)']:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)

        df_reprobados = df[df[p['TOTAL']] < p['UMBRAL']].copy()
        
        estudiantes_listos = []
        for _, row in df_reprobados.iterrows():
            # Limpieza de nombres para asegurar que no salgan como "0.0"
            nom = str(row['Nombre']).strip() if pd.notna(row['Nombre']) else ""
            ape = str(row['Apellido(s)']).strip() if pd.notna(row['Apellido(s)']) else ""
            nombre_completo = f"{nom} {ape}".strip()

            if parcial_id == '1':
                estrategia = "Acompañamiento pedagógico continuo y tutoría académica preventiva."
                espacio = "Aula de clase / Tutorías individuales."
                semana = str(random.randint(2, 7))
            else:
                estrategia = "Recuperación de aprendizajes mediante tutorías grupales e individuales."
                espacio = "Escenarios de tutoría (Presencial/Virtual)."
                semana = "17"

            estudiantes_listos.append({
                "correlacional": len(estudiantes_listos) + 1,
                "nombre": nombre_completo,
                "calificacion": float(row[p['TOTAL']]),
                "estrategia": estrategia,
                "espacio": espacio,
                "semana": semana
            })
        
        return estudiantes_listos, p

    def _generate_report(self, datos_generales, estudiantes, info_parcial):
        template_path = self.get_template_path()
        document = Document(template_path)

        replacements = {f"{{{{{self.PLACEHOLDER_MAP[k]}}}}}": v for k, v in datos_generales.items() if k in self.PLACEHOLDER_MAP}
        
        # CORRECCIÓN: Ahora coincide con tu Word {{parcial_texto}}
        replacements["{{parcial_texto}}"] = info_parcial['TEXTO']
        replacements["{{nota_minima}}"] = f"{info_parcial['UMBRAL']:.2f}"

        for paragraph in document.paragraphs:
            self._replace_placeholders_in_paragraph(paragraph, replacements)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(paragraph, replacements)

        # TABLA 1: Estudiantes (Nombre Completo)
        if len(document.tables) >= 1:
            t1 = document.tables[0]
            for est in estudiantes:
                row = t1.add_row().cells
                row[0].text = str(est["correlacional"])
                row[1].text = est["nombre"]
                row[2].text = f"{est['calificacion']:.2f}"

        # TABLA 2: Plan de Acción (Solo Nro Correlacional en columna Estudiante)
        if len(document.tables) >= 2:
            t2 = document.tables[1]
            for est in estudiantes:
                row = t2.add_row().cells
                # CORRECCIÓN: En la columna Estudiante (index 0) solo va el Nro Correlacional
                row[0].text = str(est["correlacional"])
                row[3].text = est["estrategia"]
                row[4].text = est["espacio"]
                row[5].text = est["semana"]

        output = BytesIO()
        document.save(output)
        output.seek(0)
        return output

app = Flask(__name__)
logic = ReportGeneratorLogic()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_report', methods=['POST'])
def generate_report():
    try:
        parcial_id = request.form.get('parcial_seleccionado')
        datos = {k: request.form.get(k) for k in request.form.keys()}
        file = request.files.get('excel_file')
        
        if not file: return "Archivo faltante", 400

        asignatura = logic._extract_subject_from_filename(file.filename)
        datos["asignatura"] = asignatura

        estudiantes, info_p = logic._process_excel_data(file, parcial_id)
        
        if not estudiantes:
            return f"Sin estudiantes con nota menor a {info_p['UMBRAL']}.", 200

        docx = logic._generate_report(datos, estudiantes, info_p)
        
        filename = f"Informe_{info_p['TEXTO']}_PARCIAL.docx"
        return send_file(docx, as_attachment=True, download_name=filename, 
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return render_template('index.html', error=str(e), data=request.form)

if __name__ == '__main__':
    app.run(debug=True)